#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# File: i_fileutils.py
# Author: Wadih Khairallah
# Description: 
# Created: 2024-12-01 12:12:08
# Modified: 2024-12-02 13:40:41

import json
import math
import re
import os
import csv
import yaml
import subprocess
import magic
import hashlib
import pytesseract
import shutil
import requests 
import pandas as pd
import speech_recognition as sr
import pdfplumber
import pathspec
from collections import Counter

from docx import Document
from datetime import datetime
from bs4 import BeautifulSoup
from io import StringIO
from PIL import Image
from xml.etree import ElementTree
from pydub import AudioSegment
from configparser import ConfigParser


#################################
# File Operation Short Cuts
#################################

# Core Utilities
def is_valid_path_pattern(string):
    """
    Checks if a given string is a valid file or directory path pattern.

    :param string: The string to check.
    :return: True if the string matches a valid path pattern, False otherwise.
    """
    if not isinstance(string, str):
        return False

    # OS-agnostic path separators
    path_sep_pattern = re.escape(os.sep)  # Current OS
    alt_sep_pattern = re.escape(os.altsep) if os.altsep else ""  # Alternative sep if exists
    path_separators = f"[{path_sep_pattern}{alt_sep_pattern}]"

    # File path regex pattern
    # Matches strings like "folder/file.txt", "C:/folder/file", "/etc/config"
    path_pattern = rf"(^\.?{path_separators}|^[a-zA-Z]:{path_separators}|^{path_separators})[^:*?\"<>|\r\n]+"

    return re.match(path_pattern, string) is not None


def guess_os_from_path(path):
    """
    Guesses the closest OS based on the path's construction.

    :param path: The path string to analyze.
    :return: A string indicating the closest OS ("Windows", "Unix", or "Unknown").
    """
    if not isinstance(path, str):
        return "Unknown"

    if re.match(r"^[a-zA-Z]:\\", path):  # Matches "C:\" style paths
        return "Windows"
    elif re.match(r"^/|~", path):  # Matches "/etc/config" or "~/config"
        return "Unix"
    else:
        return "Unknown"


def is_absolute_path(path):
    """
    Checks if a given path is an absolute path.

    :param path: The path string to check.
    :return: True if the path is absolute, False otherwise.
    """
    if not isinstance(path, str):
        return False

    return os.path.isabs(path)


def normalize_path(path):
    """
    Normalizes a path for the current OS.

    :param path: The path string to normalize.
    :return: A normalized path string.
    """
    if not isinstance(path, str):
        return None

    return os.path.normpath(path)


def clean_path(path):
    """
    Cleans and normalizes a file or directory path.

    Returns:
        - For file/directory: Normalized absolute path if it exists.
        - None if the path is invalid.
    """
    path = os.path.abspath(os.path.expanduser(path))
    return path if os.path.exists(path) else None

expand_path = clean_path


# File and Directory Operations
def move_file(source_path, destination_path):
    """
    Moves a file from source_path to destination_path.
    """
    return _move(source_path, destination_path, is_file=True)

def move_directory(source_path, destination_path):
    """
    Moves a directory from source_path to destination_path.
    """
    return _move(source_path, destination_path, is_file=False)

def _move(source_path, destination_path, is_file):
    source_path = clean_path(source_path)
    destination_path = clean_path(destination_path) or os.path.abspath(destination_path)

    if not source_path or not (os.path.isfile(source_path) if is_file else os.path.isdir(source_path)):
        print(f"Source {'file' if is_file else 'directory'} does not exist or is invalid: {source_path}")
        return False

    try:
        destination_dir = os.path.dirname(destination_path)
        os.makedirs(destination_dir, exist_ok=True)
        shutil.move(source_path, destination_path)
        print(f"{'File' if is_file else 'Directory'} moved successfully: {source_path} -> {destination_path}")
        return True
    except Exception as e:
        print(f"Error moving {'file' if is_file else 'directory'}: {e}")
        return False


def copy_file(source_path, destination_path):
    """
    Copies a file from source_path to destination_path.
    """
    return _copy(source_path, destination_path, is_file=True)


def copy_directory(source_path, destination_path):
    """
    Copies a directory from source_path to destination_path.
    """
    return _copy(source_path, destination_path, is_file=False)


def _copy(source_path, destination_path, is_file):
    source_path = clean_path(source_path)
    destination_path = clean_path(destination_path) or os.path.abspath(destination_path)

    if not source_path or not (os.path.isfile(source_path) if is_file else os.path.isdir(source_path)):
        print(f"Source {'file' if is_file else 'directory'} does not exist or is invalid: {source_path}")
        return False

    try:
        os.makedirs(os.path.dirname(destination_path), exist_ok=True)
        if is_file:
            shutil.copy2(source_path, destination_path)
        else:
            shutil.copytree(source_path, destination_path, dirs_exist_ok=True)
        print(f"{'File' if is_file else 'Directory'} copied successfully: {source_path} -> {destination_path}")
        return True
    except Exception as e:
        print(f"Error copying {'file' if is_file else 'directory'}: {e}")
        return False


def remove_file(file_path):
    """
    Removes a file.
    """
    file_path = clean_path(file_path)

    if not file_path or not os.path.isfile(file_path):
        print(f"File does not exist: {file_path}")
        return False

    try:
        os.remove(file_path)
        print(f"File removed successfully: {file_path}")
        return True
    except Exception as e:
        print(f"Error removing file: {e}")
        return False


def remove_directory(directory_path):
    """
    Removes a directory and its contents.
    """
    directory_path = clean_path(directory_path)

    if not directory_path or not os.path.isdir(directory_path):
        print(f"Directory does not exist or is invalid: {directory_path}")
        return False

    try:
        shutil.rmtree(directory_path)
        print(f"Directory and all contents removed successfully: {directory_path}")
        return True
    except Exception as e:
        print(f"Error removing directory: {e}")
        return False


def is_empty_file(file_path):
    """
    Checks if a file is empty.
    """
    file_path = clean_path(file_path)
    if not file_path or not os.path.isfile(file_path):
        print(f"File does not exist or is invalid: {file_path}")
        return False
    return os.path.getsize(file_path) == 0


def is_empty_directory(directory_path):
    """
    Checks if a directory is empty.
    """
    directory_path = clean_path(directory_path)
    if not directory_path or not os.path.isdir(directory_path):
        print(f"Directory does not exist or is invalid: {directory_path}")
        return False
    return not os.listdir(directory_path)


def get_sha256(file_path):
    """
    Computes the SHA-256 hash of a file.

    :param file_path: Path to the file.
    :return: The SHA-256 hash as a hexadecimal string.
    :raises FileNotFoundError: If the file does not exist.
    :raises IOError: If an error occurs while reading the file.
    """
    sha256_hash = hashlib.sha256()

    try:
        with open(file_path, "rb") as f:  # Open the file in binary mode
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
    except FileNotFoundError:
        print(f"File not found: {file_path}")
        raise
    except IOError as e:
        print(f"Error reading file {file_path}: {e}")
        raise

    return sha256_hash.hexdigest()


def read_text_file_content(path):
    """
    Reads the content of a file safely.

    :param path: Path to the file.
    :return: The content of the file as a string, or None if reading fails.
    """
    if not is_file(path):
        return None

    try:
        with open(path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error reading file: {e}")
        return None


def extract_exif(file_path):
    """
    Extracts EXIF metadata from a file using exiftool.

    :param file_path: Path to the file.
    :return: A dictionary of EXIF metadata, or None if extraction fails.
    """
    exif_data = None

    try:
        # Run exiftool command
        result = subprocess.run(
            ['exiftool', '-j', file_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False
        )

        # Check for success
        if result.returncode == 0:
            exif_data = json.loads(result.stdout.decode(errors='ignore'))[0]
        else:
            # Handle non-zero return codes
            error_message = result.stderr.decode(errors='ignore')
            print(f"Exiftool error (code {result.returncode}): {error_message}")

    except FileNotFoundError:
        print("Exiftool is not installed or not found in PATH.")
    except json.JSONDecodeError:
        print("Failed to parse EXIF data from exiftool output.")
    except Exception as e:
        print(f"Unexpected error during EXIF extraction: {e}")

    return exif_data


def remove_special_chars(values):
    """
    Removes special characters from a string or a collection of strings.

    :param values: A string, list, or tuple of strings.
    :return: A string or a collection of strings with special characters removed.
    """
    def clean(value):
        """Helper function to clean a single string."""
        return re.sub(r'[^\-\.,\#A-Za-z0-9 ]+', '', value)

    if isinstance(values, str):
        return clean(values)
    elif isinstance(values, (list, tuple)):
        # Clean each element in the collection
        cleaned_values = [clean(value) for value in values]
        return type(values)(cleaned_values)  # Preserve input type (list/tuple)
    else:
        raise TypeError("Input must be a string, list, or tuple.")



def get_mime_type(path):
    """
    Retrieves the MIME type of a file using both mimetypes and python-magic.
    """
    if not is_file(path):
        return None

    # Try mimetypes first
    mime_type, _ = mimetypes.guess_type(path)
    if mime_type:
        return mime_type

    # Fall back to python-magic for deeper inspection
    try:
        magic_mime = magic.Magic(mime=True)
        return magic_mime.from_file(path)
    except Exception as e:
        print(f"Error detecting MIME type: {e}")
        return None


def extract_text(file_path):
    """
    Extracts text or content from a file using appropriate processors based on its MIME type.

    :param file_path: Path to the file.
    :return: Extracted content as a string, or None if extraction fails.
    """
    file_path = clean_path(file_path)
    if not file_path:
        print(f"Invalid file path: {file_path}")
        return None

    mime_type = get_mime_type(file_path)
    if not mime_type:
        # Attempt metadata extraction as a fallback
        return FileMetadataExtractor(file_path).generate_report(output="markdown")

    try:
        # Define MIME type categories
        text_types = {
            'application/json',
            'application/xml',
            'application/x-yaml',
            'text/markdown'
        }
        spreadsheet_types = {
            'application/vnd.ms-excel',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }

        # Process based on MIME type
        if mime_type.startswith('text/') or mime_type in text_types:
            # Text-based files
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif mime_type in spreadsheet_types:
            # Spreadsheets
            return excel_to_csv(file_path)

        elif mime_type == 'application/pdf':
            # PDFs
            processor = PDFProcessor(file_path)
            return processor.to_markdown()

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Word Documents
            processor = WordProcessor(file_path)
            return processor.to_markdown()

        elif mime_type.startswith('image/'):
            # Images
            processor = ImageProcessor(file_path)
            return processor.process()

        elif mime_type.startswith('audio/'):
            # Audio
            processor = AudioProcessor(file_path)
            return processor.transcribe()

        else:
            # Catch-All for unsupported types
            processor = FileMetadataExtractor(file_path)
            return processor.generate_report(output="markdown")

    except Exception as e:
        print(f"Error processing {file_path} with MIME type {mime_type}: {e}")
        return None



def extract_dir_text(dir_path):
    """
    Extracts text or content from all files in a directory, respecting .gitignore rules.

    :param dir_path: Path to the directory.
    :return: Combined content from all files in the directory.
    """
    dir_path = clean_path(dir_path)
    if not os.path.isdir(dir_path):
        print(f"Invalid directory path: {dir_path}")
        return None

    # Check for a .gitignore file and parse it
    gitignore_path = os.path.join(dir_path, '.gitignore')
    spec = None
    if os.path.isfile(gitignore_path):
        with open(gitignore_path, 'r') as f:
            gitignore = f.read()
        spec = pathspec.PathSpec.from_lines(pathspec.patterns.GitWildMatchPattern, gitignore.splitlines())

    content = ""
    for root, dirs, files in os.walk(dir_path):
        # Skip hidden directories
        dirs[:] = [d for d in dirs if not d.startswith('.')]

        for file in files:
            # Skip hidden files
            if file.startswith('.'):
                continue

            file_path = os.path.join(root, file)
            file_path = clean_path(file_path)

            # Skip files matching .gitignore patterns
            if spec and spec.match_file(file_path):
                continue

            try:
                # Extract file content
                file_content = extract_text(file_path)
                if file_content:
                    content += f"## File: {file_path}\n"
                    content += "```\n"
                    content += file_content
                    content += "\n```\n\n"
            except Exception as e:
                print(f"Error processing file {file_path}: {e}")

    return content


def break_text(text, num_words):
    words = text.split()
    chunks = [' '.join(words[i:i + num_words]) for i in range(0, len(words), num_words)]
    return chunks

def lucene_like_to_regex(query):
    # Replace field:term to term
    single_term_regex = re.sub(r'\S+:(\S+)', r'\1', query)

    # Escape special regex characters, but leave our syntax elements
    escaped = re.sub(r'([\\.+^$[\]{}=!<>|:,\-])', r'\\\1', single_term_regex)

    # Restore escaped spaces (i.e., '\ ' to ' ')
    escaped = re.sub(r'\\ ', ' ', escaped)

    # Process grouping parentheses and quoted strings
    groups_and_quotes = re.sub(r'([()])', r'\\\1', escaped)
    groups_and_quotes = re.sub(r'"(.*?)"', r'\1', groups_and_quotes)

    # Convert wildcard queries to regex
    wildcard_regex = groups_and_quotes.replace('?', '.').replace('*', '.*')

    # Convert TO (range) queries to regex
    range_regex = re.sub(r'\[(\d+)\sTO\s(\d+)\]', lambda m: f"[{m.group(1)}-{m.group(2)}]", wildcard_regex)

    # Convert AND, OR and NOT queries to regex
    # AND operator is a bit tricky. We use positive lookaheads to emulate AND behavior in regex
    and_operator_regex = re.sub(r'(\S+)\sAND\s(\S+)', r'(?=.*\1)(?=.*\2)', range_regex)
    or_operator_regex = and_operator_regex.replace(' OR ', '|')
    not_operator_regex = or_operator_regex.replace(' NOT ', '^(?!.*')

    # Closing parentheses for each NOT operator
    final_regex = not_operator_regex.replace(' ', ').*')

    try:
        re.compile(final_regex)
        return final_regex
    except re.error:
        print(f"Invalid search term: {query}")
        return False

def downloadImage(url):
    if is_image(url):
        filename = os.path.basename(urlparse(url).path)
        save_path = os.path.join('/tmp/', filename)

        response = requests.get(url, stream=True)
        response.raise_for_status()

        with open(save_path, 'wb') as file:
            for chunk in response.iter_content(chunk_size=8192):
                file.write(chunk)

        return clean_path(save_path)
    else:
        print(f"Unable to pull image from {url}")
        return None

def is_image(file_path_or_url):
    try:
        if is_url(file_path_or_url):
            response = requests.head(file_path_or_url, allow_redirects=True)
            content_type = response.headers.get("Content-Type", "").lower()
            if content_type.startswith("image/"):
                return True
        else:
            mime = get_mime_type(file_path_or_url)
            if mime.startswith("image/"):
                return true
    except Exception as e:
        return False
G



def excel_to_csv(file_path):
    """
    Converts an Excel file to its text-based CSV representation.
    
    :param file_path: Path to the Excel file.
    :return: Text content of the file in CSV format, or None if conversion fails.
    """
    file_path = clean_path(file_path)
    if not file_path:
        print("Invalid file path.")
        return None

    try:
        # Read the Excel file into a DataFrame
        df = pd.read_excel(file_path)

        # Convert the DataFrame to CSV format as text
        output = StringIO()
        df.to_csv(output, index=False)
        csv_content = output.getvalue()
        output.close()

        return csv_content

    except Exception as e:
        print(f"Failed to convert Excel to text: {e}")
        return None


class ImageProcessor:
    """
    Processor for image files to extract metadata and text, and convert to Markdown format.
    """

    def __init__(self, file_path):
        self.file_path = clean_path(file_path)
        self.mime_type = get_mime_type(self.file_path)
        self.metadata = {}
        self.extracted_text = ""

    def process(self):
        """
        Processes the image file by extracting metadata and performing OCR.

        :return: Markdown-formatted string with image metadata and text.
        """
        if not self.file_path:
            print("Invalid file path.")
            return None

        if not self.mime_type or not self.mime_type.startswith("image/"):
            print("File is not a valid image.")
            return None

        try:
            # Open the image and extract metadata
            with Image.open(self.file_path) as img:
                self._extract_metadata(img)

                # Perform OCR to extract text
                self.extracted_text = self._perform_ocr(img)

                # Generate Markdown
                return self._generate_markdown()

        except Exception as e:
            print(f"Failed to process image: {e}")
            return None

    def _extract_metadata(self, img):
        """
        Extracts metadata from the image.

        :param img: PIL Image object.
        """
        self.metadata = {
            "format": img.format,
            "size": f"{img.size[0]}x{img.size[1]}",  # Width x Height
            "mode": img.mode,  # Color mode
            "mime_type": self.mime_type,
        }

    @staticmethod
    def _perform_ocr(img):
        """
        Performs OCR on the image to extract text.

        :param img: PIL Image object.
        :return: Extracted text as a string.
        """
        try:
            return pytesseract.image_to_string(img).strip()
        except Exception as e:
            print(f"Failed to perform OCR: {e}")
            return ""

    def _generate_markdown(self):
        """
        Generates a Markdown document with metadata and extracted text.

        :return: Markdown-formatted string.
        """
        markdown_content = StringIO()
        markdown_content.write("# Image Analysis Report\n\n")

        # Add metadata
        markdown_content.write("## Metadata\n")
        for key, value in self.metadata.items():
            markdown_content.write(f"- **{key.capitalize()}**: {value}\n")

        # Add extracted text
        markdown_content.write("\n## Extracted Text\n")
        if self.extracted_text:
            markdown_content.write("```\n")
            markdown_content.write(self.extracted_text)
            markdown_content.write("\n```\n")
        else:
            markdown_content.write("No text could be extracted from the image.\n")

        return markdown_content.getvalue()


class FileMetadataExtractor:
    """
    A catch-all file analyzer for extracting metadata and generating detailed reports.
    """

    def __init__(self, file_path):
        self.file_path = clean_path(file_path)
        self.metadata = {}

    def analyze(self):
        """
        Performs the file analysis and populates metadata.
        """
        if not self.file_path:
            raise FileNotFoundError("Invalid file path.")

        file_stats = os.stat(self.file_path)
        creation_time = datetime.fromtimestamp(file_stats.st_ctime)
        modified_time = datetime.fromtimestamp(file_stats.st_mtime)

        self.metadata = {
            "file_path": str(self.file_path),
            "file_size": str(file_stats.st_size),
            "creation_time": str(creation_time),
            "modification_time": str(modified_time),
            "permissions": oct(file_stats.st_mode & 0o777),
            "mime_type": str(get_mime_type(self.file_path, mime=True)),
            "hashes": {},
            "readable_strings": [],
            "magic_numbers": None,
            "embedded_urls": [],
            "entropy": None,
            "exif": {},
        }

        self._extract_exif()
        self._analyze_binary()

    def _extract_exif(self):
        """
        Extracts EXIF data from the file.
        """
        exif_data = extract_exif(self.file_path)
        if exif_data:
            self.metadata["exif"] = exif_data

    def _analyze_binary(self):
        """
        Analyzes the binary content of the file for hashes, entropy, strings, and URLs.
        """
        def calculate_entropy(data):
            """Calculate Shannon entropy."""
            if not data:
                return 0
            counter = Counter(data)
            length = len(data)
            return -sum((count / length) * math.log2(count / length) for count in counter.values())

        def extract_strings(data):
            """Extract readable ASCII and Unicode strings."""
            ascii_regex = re.compile(rb'[ -~]{4,}')
            unicode_regex = re.compile(rb'(?:[\x20-\x7E][\x00]){4,}')
            strings = ascii_regex.findall(data) + unicode_regex.findall(data)
            return [s.decode(errors='ignore') for s in strings]

        def find_binary_urls(data):
            """Find embedded URLs."""
            url_regex = re.compile(rb'https?://[^\s]+')
            return [url.decode(errors='ignore') for url in url_regex.findall(data)]

        try:
            with open(self.file_path, 'rb') as file:
                binary_data = file.read()

                # Compute hashes
                self.metadata["hashes"]["sha256"] = hashlib.sha256(binary_data).hexdigest()
                self.metadata["hashes"]["md5"] = hashlib.md5(binary_data).hexdigest()

                # Extract strings and URLs
                self.metadata["readable_strings"] = extract_strings(binary_data)[:100]  # Limit to 100 strings
                self.metadata["embedded_urls"] = find_binary_urls(binary_data)

                # Calculate entropy
                self.metadata["entropy"] = calculate_entropy(binary_data)

                # Extract magic numbers
                self.metadata["magic_numbers"] = binary_data[:4].hex()

        except Exception as e:
            print(f"Error analyzing binary data: {e}")

    def generate_report(self, output="json"):
        """
        Generates a detailed report in JSON or Markdown format.

        :param output: Report format ("json" or "markdown").
        :return: The report as a string.
        """
        if output == "json":
            return json.dumps(self.metadata, indent=4)
        elif output == "markdown":
            return self._generate_markdown_report()
        else:
            raise ValueError("Unsupported report format.")

    def _generate_markdown_report(self):
        """
        Generates a Markdown-formatted report.
        """
        md = f"# Catch-All Analysis Report\n\n"
        md += f"**File Path**: {self.metadata['file_path']}\n"
        md += f"**File Size**: {self.metadata['file_size']} bytes\n"
        md += f"**MIME Type**: {self.metadata['mime_type']}\n"
        md += f"**Permissions**: {self.metadata['permissions']}\n"
        md += f"**Creation Time**: {self.metadata['creation_time']}\n"
        md += f"**Modification Time**: {self.metadata['modification_time']}\n"
        md += f"**SHA-256 Hash**: {self.metadata['hashes']['sha256']}\n"
        md += f"**MD5 Hash**: {self.metadata['hashes']['md5']}\n"
        md += f"**Magic Numbers**: {self.metadata['magic_numbers']}\n"
        md += f"**Entropy**: {self.metadata['entropy']:.4f}\n\n"

        if self.metadata.get("exif"):
            md += "## Exif Data\n"
            for key, value in self.metadata["exif"].items():
                md += f"- **{key}**: {value}\n"

        md += "\n## Readable Strings\n"
        md += "\n".join(f"- {string}" for string in self.metadata["readable_strings"][:10])
        md += "\n\n## Embedded URLs\n"
        md += "\n".join(f"- {url}" for url in self.metadata["embedded_urls"])
        md += "\n"

        return md

def is_dir(path):
    """
    Checks if a given path is a directory.

    :param path: The path string to check.
    :return: True if the path is a directory, False otherwise.
    """
    path = clean_path(path)
    return os.path.isdir(path) if path else False


def is_file(path):
    """
    Checks if a given path is a file.

    :param path: The path string to check.
    :return: True if the path is a file, False otherwise.
    """
    path = clean_path(path)
    return os.path.isfile(path) if path else False


def is_text_file(path):
    """
    Checks if a file is a plain text file.
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type.startswith("text/")

def is_image_file(path):
    """
    Checks if a file is an image.
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type and mime_type.startswith("image/")

def is_audio_file(path):
    """
    Checks if a file is an audio file.
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type and mime_type.startswith("audio/")


def is_video_file(path):
    """
    Checks if a file is a video file.
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type and mime_type.startswith("video/")

def is_pdf_file(path):
    """
    Checks if a file is a PDF.
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type == "application/pdf"

def is_archive_file(path):
    """
    Checks if a file is a compressed archive (e.g., zip, tar, gz).
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type and mime_type in [
        "application/zip",
        "application/x-tar",
        "application/gzip",
        "application/x-7z-compressed",
        "application/x-rar-compressed",
    ]


def is_executable_file(path):
    """
    Checks if a file is an executable.
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type and mime_type in [
        "application/x-executable",
        "application/x-msdownload",
    ]

def is_spreadsheet_file(path):
    """
    Checks if a file is a spreadsheet (e.g., Excel, CSV).
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type and mime_type in [
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "text/csv",
    ]

def is_script_file(path):
    """
    Checks if a file is a script (e.g., Python, Bash, Perl).
    """
    if not is_file(path):
        return False

    mime_type = get_mime_type(path)
    return mime_type and mime_type in [
        "text/x-python",
        "text/x-shellscript",
        "application/x-perl",
        "application/x-javascript",
    ]

def is_json(path):
    """
    Checks if a file is a valid JSON file.

    :param path: Path to the file.
    :return: True if the file is valid JSON, False otherwise.
    """
    content = read_text_file_content(path)
    if content is None:
        return False

    try:
        json.loads(content)
        return True
    except json.JSONDecodeError:
        return False

def is_yaml(path):
    """
    Checks if a file is a valid YAML file.

    :param path: Path to the file.
    :return: True if the file is valid YAML, False otherwise.
    """
    content = read_text_file_content(path)
    if content is None:
        return False

    try:
        yaml.safe_load(content)
        return True
    except yaml.YAMLError:
        return False

def is_ini(path):
    """
    Checks if a file is a valid INI file.

    :param path: Path to the file.
    :return: True if the file is valid INI, False otherwise.
    """
    content = read_text_file_content(path)
    if content is None:
        return False

    parser = ConfigParser()
    try:
        parser.read_string(content)
        return True
    except configparser.Error:
        return False


def is_xml(path):
    """
    Checks if a file is a valid XML file.

    :param path: Path to the file.
    :return: True if the file is valid XML, False otherwise.
    """
    content = read_text_file_content(path)
    if content is None:
        return False

    try:
        ElementTree.fromstring(content)
        return True
    except ElementTree.ParseError:
        return False


def is_csv(path):
    """
    Checks if a file is a valid CSV file.

    :param path: Path to the file.
    :return: True if the file is valid CSV, False otherwise.
    """
    if not is_file(path):
        return False

    try:
        with open(path, newline='', encoding='utf-8') as file:
            csv.Sniffer().sniff(file.read(1024))
        return True
    except (csv.Error, IOError):
        return False


def is_html(path):
    """
    Checks if a file is a valid HTML file.

    :param path: Path to the file.
    :return: True if the file is valid HTML, False otherwise.
    """
    content = read_text_file_content(path)
    if content is None:
        return False

    try:
        soup = BeautifulSoup(content, "html.parser")
        return bool(soup.find())
    except Exception as e:
        print(f"Error parsing HTML: {e}")
        return False


def validate_file(path):
    """
    Validates a file and returns a dictionary of details and mismatches.

    :param path: Path to the file.
    :return: Dictionary with detailed file validation information.
    """
    if not os.path.isfile(path):
        return {"error": "Invalid file or path does not exist."}

    result = {
        "path": path,
        "exists": os.path.exists(path),
        "extension": get_extension(path),
        "mime_type": get_mime_type(path),
        "magic_type": None,
        "stat": None,
        "exif_data": None,
        "mismatches": [],
    }

    # Get MIME type using magic
    try:
        magic_mime = magic.Magic(mime=True)
        result["magic_type"] = magic_mime.from_file(path)
    except Exception as e:
        result["magic_type"] = f"Error: {e}"

    # Collect file statistics
    try:
        result["stat"] = os.stat(path)
    except Exception as e:
        result["stat"] = f"Error: {e}"

    # Get EXIF data for image files
    if result["mime_type"] and result["mime_type"].startswith("image/"):
        try:
            with Image.open(path) as img:
                result["exif_data"] = img._getexif() or {}
        except Exception as e:
            result["exif_data"] = f"Error: {e}"

    # Mismatch detection
    if result["extension"] and result["mime_type"]:
        extension_to_mime = {
            ".txt": "text/plain",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".json": "application/json",
            ".yaml": "application/x-yaml",
            ".yml": "application/x-yaml",
            ".xml": "application/xml",
            ".html": "text/html",
            ".csv": "text/csv",
            ".mp3": "audio/mpeg",
            ".mp4": "video/mp4",
            ".zip": "application/zip",
            ".pdf": "application/pdf",
        }
        expected_mime = extension_to_mime.get(result["extension"])
        if expected_mime and expected_mime != result["mime_type"]:
            result["mismatches"].append(
                {
                    "type": "Extension vs MIME type",
                    "expected": expected_mime,
                    "actual": result["mime_type"],
                }
            )

    if result["mime_type"] and result["magic_type"] and result["mime_type"] != result["magic_type"]:
        result["mismatches"].append(
            {
                "type": "MIME vs Content",
                "expected": result["mime_type"],
                "actual": result["magic_type"],
            }
        )

    return result


class PDFProcessor:
    """
    Processor for PDF files to extract content as Markdown.
    """

    def __init__(self, pdf_path):
        self.pdf_path = clean_path(pdf_path)

    def to_markdown(self):
        """
        Converts a PDF to Markdown, extracting metadata, text, tables, and images.
        """
        if not self.pdf_path:
            print("Invalid PDF path.")
            return None

        markdown_content = ""

        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                # Extract metadata
                markdown_content += self._extract_metadata(pdf)

                # Process each page
                for page_num, page in enumerate(pdf.pages, start=1):
                    markdown_content += f"\n\n## Page {page_num}\n\n"
                    markdown_content += self._safe_extract(self._extract_page_text, page, "No text found on this page.\n\n")
                    markdown_content += self._safe_extract(self._extract_page_tables, page, "No tables found on this page.\n\n")
                    markdown_content += self._safe_extract(self._extract_page_images, page, page_num, "No images found on this page.\n\n")
        except Exception as e:
            print(f"Critical error processing PDF {self.pdf_path}: {e}")
            markdown_content += f"Error processing PDF: {e}\n"

        return markdown_content

    @staticmethod
    def _safe_extract(func, *args, fallback=""):
        """
        Safely executes a function, returning a fallback message if it fails.

        :param func: The function to execute.
        :param args: Arguments for the function.
        :param fallback: Fallback message if the function fails.
        :return: The function's return value or the fallback message.
        """
        try:
            return func(*args)
        except Exception as e:
            print(f"Error in {func.__name__}: {e}")
            return fallback

    def _extract_metadata(self, pdf):
        """
        Extracts and formats PDF metadata into Markdown.

        :param pdf: A pdfplumber PDF object.
        :return: Markdown-formatted metadata.
        """
        metadata_md = "## Metadata\n"
        metadata = pdf.metadata or {}
        for key, value in metadata.items():
            metadata_md += f"- **{key}**: {value}\n"
        return metadata_md + "\n"

    def _extract_page_text(self, page):
        """
        Extracts text from a PDF page.

        :param page: A pdfplumber page object.
        :return: Extracted text as a string.
        """
        text = page.extract_text()
        return text + "\n\n" if text else "No text found on this page.\n\n"

    def _extract_page_tables(self, page):
        """
        Extracts tables from a PDF page and converts them to Markdown.

        :param page: A pdfplumber page object.
        :return: Markdown-formatted tables.
        """
        tables_md = ""
        tables = page.extract_tables() or []
        for table_num, table in enumerate(tables, start=1):
            tables_md += f"### Table {table_num}\n\n"
            tables_md += self._table_to_markdown(table)
        return tables_md if tables_md else "No tables found on this page.\n\n"

    @staticmethod
    def _table_to_markdown(table):
        """
        Converts a table into Markdown format.

        :param table: A table (list of lists).
        :return: Markdown-formatted table.
        """
        table_md = ""
        for row in table:
            table_md += "| " + " | ".join(row) + " |\n"
        if table:
            table_md += "|---" * len(table[0]) + "|\n"
        return table_md

    def _extract_page_images(self, page, page_num):
        """
        Extracts images from a PDF page and converts them to Markdown.

        :param page: A pdfplumber page object.
        :param page_num: Page number of the image.
        :return: Markdown-formatted image links with optional OCR text.
        """
        images_md = ""
        for image_num, image in enumerate(page.images, start=1):
            if "data" in image:
                image_data = image["data"]
                images_md += self._save_image(image_data, page_num, image_num)
        return images_md if images_md else "No images found on this page.\n\n"

    @staticmethod
    def _save_image(image_data, page_num, image_num, save_dir="./"):
        """
        Saves an image from a PDF and returns its Markdown representation.

        :param image_data: Binary data of the image.
        :param page_num: Page number of the image.
        :param image_num: Image index on the page.
        :param save_dir: Directory to save the image.
        :return: Markdown link to the saved image.
        """
        os.makedirs(save_dir, exist_ok=True)
        image_path = os.path.join(save_dir, f"page-{page_num}-image-{image_num}.png")
        try:
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)
            return f"![Image {image_num}](./{image_path})\n\n"
        except Exception as e:
            print(f"Error saving image: {e}")
            return f"Error saving image {image_num} on page {page_num}.\n\n"


class WordProcessor():
    """
    Processor for Word documents (.docx).
    """
    def to_markdown(self):
        if not self.file_path:
            print("Invalid file path.")
            return None

        try:
            # Load the Word document
            doc = Document(self.file_path)
        except Exception as e:
            print(f"Error loading Word document {self.file_path}: {e}")
            return None

        markdown_content = self._process_paragraphs(doc)

        # Add tables
        try:
            for table in doc.tables:
                markdown_content += "\n" + self._table_to_markdown(table) + "\n"
        except Exception as e:
            print(f"Error processing tables in {self.file_path}: {e}")

        return markdown_content

    def _process_paragraphs(self, doc):
        markdown_content = ""
        for paragraph in doc.paragraphs:
            # Handle headings
            if paragraph.style.name.startswith("Heading"):
                heading_level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                markdown_content += f"{'#' * heading_level} {paragraph.text.strip()}\n\n"
            # Handle lists
            elif paragraph.style.name.startswith("List"):
                markdown_content += self._list_to_markdown(paragraph)
            # Handle normal text
            elif paragraph.text.strip():
                markdown_content += f"{paragraph.text.strip()}\n\n"
        return markdown_content

    def _list_to_markdown(self, paragraph):
        text = paragraph.text.strip()
        indent_level = paragraph._p.getparent().xpath('count(preceding-sibling::w:pStyle)') or 0
        if paragraph.style.name.startswith("List Bullet"):
            return f"{'  ' * int(indent_level)}- {text}\n"
        elif paragraph.style.name.startswith("List Number"):
            return f"{'  ' * int(indent_level)}1. {text}\n"
        return ""

    def _table_to_markdown(self, table):
        table_md = ""
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_md += "| " + " | ".join(cells) + " |\n"
        if len(table.rows) > 0:
            table_md += "|---" * len(table.rows[0].cells) + "|\n"
        return table_md


class AudioProcessor:
    """
    A processor for audio files, providing conversion, transcription, and other utilities.
    """

    SUPPORTED_FORMATS = ['wav', 'mp3', 'ogg', 'flac', 'aac', 'wma']

    def __init__(self, audio_file):
        self.audio_file = clean_path(audio_file)
        self.recognizer = sr.Recognizer()

    def transcribe(self, engine="google"):
        """
        Transcribes the audio file to text using the specified engine.

        :param engine: The transcription engine to use (default: "google").
        :return: Transcribed text as a string, or None if transcription fails.
        """
        audio_file = self._convert_to_wav(self.audio_file)

        if not audio_file:
            print("Failed to convert audio file to WAV format.")
            return None

        try:
            with sr.AudioFile(audio_file) as source:
                audio = self.recognizer.record(source)

            if engine.lower() == "google":
                return self.recognizer.recognize_google(audio)
            elif engine.lower() == "sphinx":
                return self.recognizer.recognize_sphinx(audio)
            else:
                print(f"Unsupported transcription engine: {engine}")
                return None

        except sr.UnknownValueError:
            print("Speech recognition could not understand audio.")
            return None
        except sr.RequestError as e:
            print(f"Could not request results from the recognition service; {e}")
            return None
        except Exception as e:
            print(f"Unexpected error during transcription: {e}")
            return None

    def convert(self, output_format="wav"):
        """
        Converts the audio file to a specified format.

        :param output_format: The desired output format (e.g., "mp3", "wav").
        :return: Path to the converted audio file, or None if conversion fails.
        """
        if output_format.lower() not in self.SUPPORTED_FORMATS:
            print(f"Unsupported format: {output_format}")
            return None

        try:
            output_file = os.path.splitext(self.audio_file)[0] + f".{output_format.lower()}"
            audio = AudioSegment.from_file(self.audio_file)
            audio.export(output_file, format=output_format)
            return output_file
        except Exception as e:
            print(f"Error converting audio file to {output_format}: {e}")
            return None

    def split_audio(self, chunk_length_ms):
        """
        Splits the audio file into smaller chunks of a specified length.

        :param chunk_length_ms: Length of each chunk in milliseconds.
        :return: List of paths to the split audio files.
        """
        try:
            audio = AudioSegment.from_file(self.audio_file)
            chunks = [
                audio[start:start + chunk_length_ms]
                for start in range(0, len(audio), chunk_length_ms)
            ]

            chunk_paths = []
            for i, chunk in enumerate(chunks):
                chunk_path = os.path.splitext(self.audio_file)[0] + f"_chunk_{i + 1}.wav"
                chunk.export(chunk_path, format="wav")
                chunk_paths.append(chunk_path)

            return chunk_paths
        except Exception as e:
            print(f"Error splitting audio file: {e}")
            return []

    def _convert_to_wav(self, audio_file):
        """
        Converts an audio file to WAV format.

        :param audio_file: Path to the input audio file.
        :return: Path to the converted WAV file, or None if conversion fails.
        """
        _, ext = os.path.splitext(audio_file)

        if ext.lower() in ['.wav', '.wave']:
            return audio_file

        try:
            wav_file = os.path.splitext(audio_file)[0] + ".wav"
            audio = AudioSegment.from_file(audio_file)
            audio.export(wav_file, format="wav")
            return wav_file
        except Exception as e:
            print(f"Error converting audio file to WAV: {e}")
            return None


