#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# File: fileutils.py
# Author: Wadih Khairallah
# Description: 
# Created: 2024-12-01 12:12:08
# Modified: 2024-12-03 00:22:51

import json
import math
import re
import os
import csv
import yaml
import subprocess
import magic
import hashlib
import pytesseract
import shutil
import requests 
import pandas as pd
import speech_recognition as sr
import pdfplumber
import pathspec
from collections import Counter

from docx import Document
from datetime import datetime
from bs4 import BeautifulSoup
from io import StringIO
from PIL import Image
from xml.etree import ElementTree
from pydub import AudioSegment
from configparser import ConfigParser

import .fsutils as fsutil
import .textutils as textutil

def extract_text(file_path):
    """
    Extracts text or content from a file using appropriate processors based on its MIME type.

    :param file_path: Path to the file.
    :return: Extracted content as a string, or None if extraction fails.
    """
    file_path = clean_path(file_path)
    if not file_path:
        print(f"Invalid file path: {file_path}")
        return None

    mime_type = get_mime_type(file_path)
    if not mime_type:
        # Attempt metadata extraction as a fallback
        return MetadataProcessor(file_path).generate_report(output="markdown")

    try:
        # Define MIME type categories
        text_types = {
            'application/json',
            'application/xml',
            'application/x-yaml',
            'text/markdown'
        }
        spreadsheet_types = {
            'application/vnd.ms-excel',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }

        # Process based on MIME type
        if mime_type.startswith('text/') or mime_type in text_types:
            # Text-based files
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif mime_type in spreadsheet_types:
            # Spreadsheets
            return excel_to_csv(file_path)

        elif mime_type == 'application/pdf':
            # PDFs
            processor = PDFProcessor(file_path)
            return processor.to_markdown()

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Word Documents
            processor = WordProcessor(file_path)
            return processor.to_markdown()

        elif mime_type.startswith('image/'):
            # Images
            processor = ImageProcessor(file_path)
            return processor.process()

        elif mime_type.startswith('audio/'):
            # Audio
            processor = AudioProcessor(file_path)
            return processor.transcribe()

        else:
            # Catch-All for unsupported types
            processor = MetadataProcessor(file_path)
            return processor.generate_report(output="markdown")

    except Exception as e:
        print(f"Error processing {file_path} with MIME type {mime_type}: {e}")
        return None


def extract_dir_text(dir_path):
    """
    Extracts text or content from all files in a directory, respecting .gitignore rules.

    :param dir_path: Path to the directory.
    :return: Combined content from all files in the directory.
    """
    dir_path = clean_path(dir_path)
    if not os.path.isdir(dir_path):
        print(f"Invalid directory path: {dir_path}")
        return None

    # Check for a .gitignore file and parse it
    gitignore_path = os.path.join(dir_path, '.gitignore')
    spec = None
    if os.path.isfile(gitignore_path):
        with open(gitignore_path, 'r') as f:
            gitignore = f.read()
        spec = pathspec.PathSpec.from_lines(pathspec.patterns.GitWildMatchPattern, gitignore.splitlines())

    content = ""
    for root, dirs, files in os.walk(dir_path):
        # Skip hidden directories
        dirs[:] = [d for d in dirs if not d.startswith('.')]

        for file in files:
            # Skip hidden files
            if file.startswith('.'):
                continue

            file_path = os.path.join(root, file)
            file_path = clean_path(file_path)

            # Skip files matching .gitignore patterns
            if spec and spec.match_file(file_path):
                continue

            try:
                # Extract file content
                file_content = extract_text(file_path)
                if file_content:
                    content += f"## File: {file_path}\n"
                    content += "```\n"
                    content += file_content
                    content += "\n```\n\n"
            except Exception as e:
                print(f"Error processing file {file_path}: {e}")

    return content


def lucene_like_to_regex(query):
    """
    Converts a Lucene-like query into a regular expression.

    :param query: Lucene-like query string.
    :return: A compiled regular expression object, or raises ValueError for invalid syntax.
    """
    def escape_lucene_syntax(text):
        """Escapes regex special characters except Lucene operators."""
        return re.sub(r'([.+^$[\]{}=!<>|\\-])', r'\\\1', text)

    try:
        # 1. Handle field:term syntax (convert `field:term` to `term`)
        query = re.sub(r'\b\w+:(\S+)', r'\1', query)

        # 2. Escape regex special characters (preserve Lucene wildcards and syntax)
        query = escape_lucene_syntax(query)

        # 3. Handle grouping parentheses and quoted strings
        query = re.sub(r'([()])', r'\\\1', query)  # Escape parentheses
        query = re.sub(r'"([^"]+)"', r'(\1)', query)  # Convert quoted strings to groups

        # 4. Convert wildcard syntax
        query = query.replace('?', '.').replace('*', '.*')

        # 5. Handle range queries ([A TO B] → [A-B])
        query = re.sub(r'\[(\w+)\sTO\s(\w+)\]', r'[\1-\2]', query)

        # 6. Convert boolean operators to regex constructs
        # AND → lookahead for both terms
        query = re.sub(r'(\S+)\sAND\s(\S+)', r'(?=.*\1)(?=.*\2)', query)

        # OR → regex alternation
        query = query.replace(' OR ', '|')

        # NOT → negative lookahead
        query = query.replace(' NOT ', '^(?!.*') + ').*'

        # Ensure the final regex is valid
        regex = re.compile(query)
        return regex

    except re.error as e:
        raise ValueError(f"Invalid Lucene-like query: {query}. Regex error: {e}")


def lucene_like_to_regex_fixed(query):
    """
    Converts a Lucene-like query into a regular expression.
    """
    def escape_lucene_syntax(text):
        """Escapes regex special characters except Lucene operators."""
        return re.sub(r'([.+^$[\]{}=!<>|\\-])', r'\\\1', text)

    try:
        # 1. Handle field:term syntax (convert `field:term` to `term`)
        query = re.sub(r'\b\w+:(\S+)', r'\1', query)

        # 2. Escape regex special characters (preserve Lucene wildcards and syntax)
        query = escape_lucene_syntax(query)

        # 3. Handle grouping parentheses and quoted strings
        query = re.sub(r'([()])', r'\\\1', query)  # Escape parentheses
        query = re.sub(r'"([^"]+)"', r'(\1)', query)  # Convert quoted strings to groups

        # 4. Convert wildcard syntax
        query = query.replace('?', '.').replace('*', '.*')

        # 5. Handle range queries ([A TO B] → [A-B])
        query = re.sub(r'\[(\w+)\sTO\s(\w+)\]', r'[\1-\2]', query)

        # 6. Convert boolean operators to regex constructs
        # AND → lookahead for both terms
        query = re.sub(r'(\S+)\sAND\s(\S+)', r'(?=.*\1)(?=.*\2)', query)

        # OR → regex alternation
        query = query.replace(' OR ', '|')

        # NOT → negative lookahead
        if 'NOT ' in query:
            query = re.sub(r'NOT (\S+)', r'^(?!.*\1).*', query)

        # Ensure the final regex is valid
        regex = re.compile(query)
        return regex

    except re.error as e:
        raise ValueError(f"Invalid Lucene-like query: {query}. Regex error: {e}")

    """
    # Retest the same queries with the fixed function
    results_fixed = {}

    for query in queries:
        try:
            regex = lucene_like_to_regex_fixed(query)
            results_fixed[query] = regex.pattern
        except ValueError as e:
            results_fixed[query] = str(e)

    results_fixed
    """


def excel_to_csv(file_path):
    """
    Converts an Excel file to its text-based CSV representation.
    
    :param file_path: Path to the Excel file.
    :return: Text content of the file in CSV format, or None if conversion fails.
    """
    file_path = clean_path(file_path)
    if not file_path:
        print("Invalid file path.")
        return None

    try:
        # Read the Excel file into a DataFrame
        df = pd.read_excel(file_path)

        # Convert the DataFrame to CSV format as text
        output = StringIO()
        df.to_csv(output, index=False)
        csv_content = output.getvalue()
        output.close()

        return csv_content

    except Exception as e:
        print(f"Failed to convert Excel to text: {e}")
        return None


class ImageProcessor:
    """
    Processor for image files to extract metadata and text, and convert to Markdown format.
    """

    def __init__(self, file_path):
        self.file_path = clean_path(file_path)
        self.mime_type = get_mime_type(self.file_path)
        self.metadata = {}
        self.extracted_text = ""

    def process(self):
        """
        Processes the image file by extracting metadata and performing OCR.

        :return: Markdown-formatted string with image metadata and text.
        """
        if not self.file_path:
            print("Invalid file path.")
            return None

        if not self.mime_type or not self.mime_type.startswith("image/"):
            print("File is not a valid image.")
            return None

        try:
            # Open the image and extract metadata
            with Image.open(self.file_path) as img:
                self._extract_metadata(img)

                # Perform OCR to extract text
                self.extracted_text = self._perform_ocr(img)

                # Generate Markdown
                return self._generate_markdown()

        except Exception as e:
            print(f"Failed to process image: {e}")
            return None

    def _extract_metadata(self, img):
        """
        Extracts metadata from the image.

        :param img: PIL Image object.
        """
        self.metadata = {
            "format": img.format,
            "size": f"{img.size[0]}x{img.size[1]}",  # Width x Height
            "mode": img.mode,  # Color mode
            "mime_type": self.mime_type,
        }

    @staticmethod
    def _perform_ocr(img):
        """
        Performs OCR on the image to extract text.

        :param img: PIL Image object.
        :return: Extracted text as a string.
        """
        try:
            return pytesseract.image_to_string(img).strip()
        except Exception as e:
            print(f"Failed to perform OCR: {e}")
            return ""

    def _generate_markdown(self):
        """
        Generates a Markdown document with metadata and extracted text.

        :return: Markdown-formatted string.
        """
        markdown_content = StringIO()
        markdown_content.write("# Image Analysis Report\n\n")

        # Add metadata
        markdown_content.write("## Metadata\n")
        for key, value in self.metadata.items():
            markdown_content.write(f"- **{key.capitalize()}**: {value}\n")

        # Add extracted text
        markdown_content.write("\n## Extracted Text\n")
        if self.extracted_text:
            markdown_content.write("```\n")
            markdown_content.write(self.extracted_text)
            markdown_content.write("\n```\n")
        else:
            markdown_content.write("No text could be extracted from the image.\n")

        return markdown_content.getvalue()


class MetadataProcessor:
    """
    A catch-all file analyzer for extracting metadata and generating detailed reports.
    """

    def __init__(self, file_path):
        self.file_path = clean_path(file_path)
        self.metadata = {}

    def analyze(self):
        """
        Performs the file analysis and populates metadata.
        """
        if not self.file_path:
            raise FileNotFoundError("Invalid file path.")

        file_stats = os.stat(self.file_path)
        creation_time = datetime.fromtimestamp(file_stats.st_ctime)
        modified_time = datetime.fromtimestamp(file_stats.st_mtime)

        self.metadata = {
            "file_path": str(self.file_path),
            "file_size": str(file_stats.st_size),
            "creation_time": str(creation_time),
            "modification_time": str(modified_time),
            "permissions": oct(file_stats.st_mode & 0o777),
            "mime_type": str(get_mime_type(self.file_path, mime=True)),
            "hashes": {},
            "readable_strings": [],
            "magic_numbers": None,
            "embedded_urls": [],
            "entropy": None,
            "exif": {},
        }

        self._extract_exif()
        self._analyze_binary()

    def _extract_exif(self):
        """
        Extracts EXIF data from the file.
        """
        exif_data = extract_exif(self.file_path)
        if exif_data:
            self.metadata["exif"] = exif_data

    def _analyze_binary(self):
        """
        Analyzes the binary content of the file for hashes, entropy, strings, and URLs.
        """
        def calculate_entropy(data):
            """Calculate Shannon entropy."""
            if not data:
                return 0
            counter = Counter(data)
            length = len(data)
            return -sum((count / length) * math.log2(count / length) for count in counter.values())

        def extract_strings(data):
            """Extract readable ASCII and Unicode strings."""
            ascii_regex = re.compile(rb'[ -~]{4,}')
            unicode_regex = re.compile(rb'(?:[\x20-\x7E][\x00]){4,}')
            strings = ascii_regex.findall(data) + unicode_regex.findall(data)
            return [s.decode(errors='ignore') for s in strings]

        def find_binary_urls(data):
            """Find embedded URLs."""
            url_regex = re.compile(rb'https?://[^\s]+')
            return [url.decode(errors='ignore') for url in url_regex.findall(data)]

        try:
            with open(self.file_path, 'rb') as file:
                binary_data = file.read()

                # Compute hashes
                self.metadata["hashes"]["sha256"] = hashlib.sha256(binary_data).hexdigest()
                self.metadata["hashes"]["md5"] = hashlib.md5(binary_data).hexdigest()

                # Extract strings and URLs
                self.metadata["readable_strings"] = extract_strings(binary_data)[:100]  # Limit to 100 strings
                self.metadata["embedded_urls"] = find_binary_urls(binary_data)

                # Calculate entropy
                self.metadata["entropy"] = calculate_entropy(binary_data)

                # Extract magic numbers
                self.metadata["magic_numbers"] = binary_data[:4].hex()

        except Exception as e:
            print(f"Error analyzing binary data: {e}")

    def generate_report(self, output="json"):
        """
        Generates a detailed report in JSON or Markdown format.

        :param output: Report format ("json" or "markdown").
        :return: The report as a string.
        """
        if output == "json":
            return json.dumps(self.metadata, indent=4)
        elif output == "markdown":
            return self._generate_markdown_report()
        else:
            raise ValueError("Unsupported report format.")

    def _generate_markdown_report(self):
        """
        Generates a Markdown-formatted report.
        """
        md = f"# Catch-All Analysis Report\n\n"
        md += f"**File Path**: {self.metadata['file_path']}\n"
        md += f"**File Size**: {self.metadata['file_size']} bytes\n"
        md += f"**MIME Type**: {self.metadata['mime_type']}\n"
        md += f"**Permissions**: {self.metadata['permissions']}\n"
        md += f"**Creation Time**: {self.metadata['creation_time']}\n"
        md += f"**Modification Time**: {self.metadata['modification_time']}\n"
        md += f"**SHA-256 Hash**: {self.metadata['hashes']['sha256']}\n"
        md += f"**MD5 Hash**: {self.metadata['hashes']['md5']}\n"
        md += f"**Magic Numbers**: {self.metadata['magic_numbers']}\n"
        md += f"**Entropy**: {self.metadata['entropy']:.4f}\n\n"

        if self.metadata.get("exif"):
            md += "## Exif Data\n"
            for key, value in self.metadata["exif"].items():
                md += f"- **{key}**: {value}\n"

        md += "\n## Readable Strings\n"
        md += "\n".join(f"- {string}" for string in self.metadata["readable_strings"][:10])
        md += "\n\n## Embedded URLs\n"
        md += "\n".join(f"- {url}" for url in self.metadata["embedded_urls"])
        md += "\n"

        return md


class PDFProcessor:
    """
    Processor for PDF files to extract content as Markdown.
    """

    def __init__(self, pdf_path):
        self.pdf_path = clean_path(pdf_path)

    def to_markdown(self):
        """
        Converts a PDF to Markdown, extracting metadata, text, tables, and images.
        """
        if not self.pdf_path:
            print("Invalid PDF path.")
            return None

        markdown_content = ""

        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                # Extract metadata
                markdown_content += self._extract_metadata(pdf)

                # Process each page
                for page_num, page in enumerate(pdf.pages, start=1):
                    markdown_content += f"\n\n## Page {page_num}\n\n"
                    markdown_content += self._safe_extract(self._extract_page_text, page, "No text found on this page.\n\n")
                    markdown_content += self._safe_extract(self._extract_page_tables, page, "No tables found on this page.\n\n")
                    markdown_content += self._safe_extract(self._extract_page_images, page, page_num, "No images found on this page.\n\n")
        except Exception as e:
            print(f"Critical error processing PDF {self.pdf_path}: {e}")
            markdown_content += f"Error processing PDF: {e}\n"

        return markdown_content

    @staticmethod
    def _safe_extract(func, *args, fallback=""):
        """
        Safely executes a function, returning a fallback message if it fails.

        :param func: The function to execute.
        :param args: Arguments for the function.
        :param fallback: Fallback message if the function fails.
        :return: The function's return value or the fallback message.
        """
        try:
            return func(*args)
        except Exception as e:
            print(f"Error in {func.__name__}: {e}")
            return fallback

    def _extract_metadata(self, pdf):
        """
        Extracts and formats PDF metadata into Markdown.

        :param pdf: A pdfplumber PDF object.
        :return: Markdown-formatted metadata.
        """
        metadata_md = "## Metadata\n"
        metadata = pdf.metadata or {}
        for key, value in metadata.items():
            metadata_md += f"- **{key}**: {value}\n"
        return metadata_md + "\n"

    def _extract_page_text(self, page):
        """
        Extracts text from a PDF page.

        :param page: A pdfplumber page object.
        :return: Extracted text as a string.
        """
        text = page.extract_text()
        return text + "\n\n" if text else "No text found on this page.\n\n"

    def _extract_page_tables(self, page):
        """
        Extracts tables from a PDF page and converts them to Markdown.

        :param page: A pdfplumber page object.
        :return: Markdown-formatted tables.
        """
        tables_md = ""
        tables = page.extract_tables() or []
        for table_num, table in enumerate(tables, start=1):
            tables_md += f"### Table {table_num}\n\n"
            tables_md += self._table_to_markdown(table)
        return tables_md if tables_md else "No tables found on this page.\n\n"

    @staticmethod
    def _table_to_markdown(table):
        """
        Converts a table into Markdown format.

        :param table: A table (list of lists).
        :return: Markdown-formatted table.
        """
        table_md = ""
        for row in table:
            table_md += "| " + " | ".join(row) + " |\n"
        if table:
            table_md += "|---" * len(table[0]) + "|\n"
        return table_md

    def _extract_page_images(self, page, page_num):
        """
        Extracts images from a PDF page and converts them to Markdown.

        :param page: A pdfplumber page object.
        :param page_num: Page number of the image.
        :return: Markdown-formatted image links with optional OCR text.
        """
        images_md = ""
        for image_num, image in enumerate(page.images, start=1):
            if "data" in image:
                image_data = image["data"]
                images_md += self._save_image(image_data, page_num, image_num)
        return images_md if images_md else "No images found on this page.\n\n"

    @staticmethod
    def _save_image(image_data, page_num, image_num, save_dir="./"):
        """
        Saves an image from a PDF and returns its Markdown representation.

        :param image_data: Binary data of the image.
        :param page_num: Page number of the image.
        :param image_num: Image index on the page.
        :param save_dir: Directory to save the image.
        :return: Markdown link to the saved image.
        """
        os.makedirs(save_dir, exist_ok=True)
        image_path = os.path.join(save_dir, f"page-{page_num}-image-{image_num}.png")
        try:
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)
            return f"![Image {image_num}](./{image_path})\n\n"
        except Exception as e:
            print(f"Error saving image: {e}")
            return f"Error saving image {image_num} on page {page_num}.\n\n"


class WordProcessor():
    """
    Processor for Word documents (.docx).
    """
    def to_markdown(self):
        if not self.file_path:
            print("Invalid file path.")
            return None

        try:
            # Load the Word document
            doc = Document(self.file_path)
        except Exception as e:
            print(f"Error loading Word document {self.file_path}: {e}")
            return None

        markdown_content = self._process_paragraphs(doc)

        # Add tables
        try:
            for table in doc.tables:
                markdown_content += "\n" + self._table_to_markdown(table) + "\n"
        except Exception as e:
            print(f"Error processing tables in {self.file_path}: {e}")

        return markdown_content

    def _process_paragraphs(self, doc):
        markdown_content = ""
        for paragraph in doc.paragraphs:
            # Handle headings
            if paragraph.style.name.startswith("Heading"):
                heading_level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                markdown_content += f"{'#' * heading_level} {paragraph.text.strip()}\n\n"
            # Handle lists
            elif paragraph.style.name.startswith("List"):
                markdown_content += self._list_to_markdown(paragraph)
            # Handle normal text
            elif paragraph.text.strip():
                markdown_content += f"{paragraph.text.strip()}\n\n"
        return markdown_content

    def _list_to_markdown(self, paragraph):
        text = paragraph.text.strip()
        indent_level = paragraph._p.getparent().xpath('count(preceding-sibling::w:pStyle)') or 0
        if paragraph.style.name.startswith("List Bullet"):
            return f"{'  ' * int(indent_level)}- {text}\n"
        elif paragraph.style.name.startswith("List Number"):
            return f"{'  ' * int(indent_level)}1. {text}\n"
        return ""

    def _table_to_markdown(self, table):
        table_md = ""
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_md += "| " + " | ".join(cells) + " |\n"
        if len(table.rows) > 0:
            table_md += "|---" * len(table.rows[0].cells) + "|\n"
        return table_md


class AudioProcessor:
    """
    A processor for audio files, providing conversion, transcription, and other utilities.
    """

    SUPPORTED_FORMATS = ['wav', 'mp3', 'ogg', 'flac', 'aac', 'wma']

    def __init__(self, audio_file):
        self.audio_file = clean_path(audio_file)
        self.recognizer = sr.Recognizer()

    def transcribe(self, engine="google"):
        """
        Transcribes the audio file to text using the specified engine.

        :param engine: The transcription engine to use (default: "google").
        :return: Transcribed text as a string, or None if transcription fails.
        """
        audio_file = self._convert_to_wav(self.audio_file)

        if not audio_file:
            print("Failed to convert audio file to WAV format.")
            return None

        try:
            with sr.AudioFile(audio_file) as source:
                audio = self.recognizer.record(source)

            if engine.lower() == "google":
                return self.recognizer.recognize_google(audio)
            elif engine.lower() == "sphinx":
                return self.recognizer.recognize_sphinx(audio)
            else:
                print(f"Unsupported transcription engine: {engine}")
                return None

        except sr.UnknownValueError:
            print("Speech recognition could not understand audio.")
            return None
        except sr.RequestError as e:
            print(f"Could not request results from the recognition service; {e}")
            return None
        except Exception as e:
            print(f"Unexpected error during transcription: {e}")
            return None

    def convert(self, output_format="wav"):
        """
        Converts the audio file to a specified format.

        :param output_format: The desired output format (e.g., "mp3", "wav").
        :return: Path to the converted audio file, or None if conversion fails.
        """
        if output_format.lower() not in self.SUPPORTED_FORMATS:
            print(f"Unsupported format: {output_format}")
            return None

        try:
            output_file = os.path.splitext(self.audio_file)[0] + f".{output_format.lower()}"
            audio = AudioSegment.from_file(self.audio_file)
            audio.export(output_file, format=output_format)
            return output_file
        except Exception as e:
            print(f"Error converting audio file to {output_format}: {e}")
            return None

    def split_audio(self, chunk_length_ms):
        """
        Splits the audio file into smaller chunks of a specified length.

        :param chunk_length_ms: Length of each chunk in milliseconds.
        :return: List of paths to the split audio files.
        """
        try:
            audio = AudioSegment.from_file(self.audio_file)
            chunks = [
                audio[start:start + chunk_length_ms]
                for start in range(0, len(audio), chunk_length_ms)
            ]

            chunk_paths = []
            for i, chunk in enumerate(chunks):
                chunk_path = os.path.splitext(self.audio_file)[0] + f"_chunk_{i + 1}.wav"
                chunk.export(chunk_path, format="wav")
                chunk_paths.append(chunk_path)

            return chunk_paths
        except Exception as e:
            print(f"Error splitting audio file: {e}")
            return []

    def _convert_to_wav(self, audio_file):
        """
        Converts an audio file to WAV format.

        :param audio_file: Path to the input audio file.
        :return: Path to the converted WAV file, or None if conversion fails.
        """
        _, ext = os.path.splitext(audio_file)

        if ext.lower() in ['.wav', '.wave']:
            return audio_file

        try:
            wav_file = os.path.splitext(audio_file)[0] + ".wav"
            audio = AudioSegment.from_file(audio_file)
            audio.export(wav_file, format="wav")
            return wav_file
        except Exception as e:
            print(f"Error converting audio file to WAV: {e}")
            return None


