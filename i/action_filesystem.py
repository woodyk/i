#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# File: action_filesystem.py
# Author: Wadih Khairallah
# Description: 
# Created: 2024-12-01 20:27:05
# Modified: 2024-12-01 22:34:01

from io import StringIO
import pandas as pd
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import speech_recognition as sr
import magic
import os
import speech_recognition as sr


def process_file(file_path: str, action: str):
    """
    Processes a file based on the specified action using FileProcessor.
    """
    try:
        if action == "summarize":
            mime_type = magic.from_file(file_path, mime=True)

            if mime_type.startswith("text/"):
                content = FileProcessor.read_text(file_path)
            elif mime_type == "application/pdf":
                content = FileProcessor.process_pdf(file_path)
            elif mime_type in [
                "application/vnd.ms-excel",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ]:
                content = FileProcessor.process_excel(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                content = FileProcessor.process_word(file_path)
            elif mime_type.startswith("image/"):
                content = FileProcessor.handle_image_file(file_path)
            elif mime_type.startswith("audio/"):
                content = FileProcessor.transcribe_audio(file_path)
            else:
                content = f"Unsupported file type: {mime_type}"

            Formatter.print(f"Processed Content:\n{content}")
        else:
            Formatter.error(f"Unknown action: {action}")
    except Exception as e:
        Formatter.error(f"Failed to process file: {e}")


def clean_path(path):
    """
    Cleans and validates a file path.
    Expands user shortcuts (~), converts to absolute path, and validates.
    """
    path = os.path.expanduser(path)
    path = os.path.abspath(path)
    if os.path.isfile(path):
        return path
    raise ValueError(f"Invalid path: {path}")

def excel_to_csv(file_path):
    """
    Converts an Excel file to CSV format in-memory and returns the content.
    """
    file_path = FileProcessor.clean_path(file_path)
    csv_content = ""
    try:
        df = pd.read_excel(file_path)
        output = StringIO()
        df.to_csv(output, index=False)
        csv_content = output.getvalue()
    except Exception as e:
        print(f"Failed to convert Excel to CSV: {e}")
    return csv_content

def process_pdf(file_path):
    """
    Processes a PDF file and extracts text.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        return text if text else "No text found in the PDF."
    except Exception as e:
        return f"Error processing PDF: {e}"


def pdf_to_markdown(file_path):
    """
    Converts a PDF file to Markdown format with text, tables, and images.
    """
    markdown_content = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                markdown_content += f"\n\n## Page {page_num}\n\n"
                text = page.extract_text()
                markdown_content += text + "\n" if text else "No text found.\n"
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables, start=1):
                        markdown_content += f"### Table {table_num}\n\n"
                        markdown_content += FileProcessor.table_to_markdown(tables)
    except Exception as e:
        print(f"Error processing PDF: {e}")
    return markdown_content

def word_to_markdown(file_path):
    """
    Converts a Word document (.docx) to Markdown format.
    Includes headings, text, tables, and lists.
    """
    file_path = FileProcessor.clean_path(file_path)
    markdown_content = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                heading_level = int(paragraph.style.name[-1])
                markdown_content += f"{'#' * heading_level} {paragraph.text.strip()}\n\n"
            elif paragraph.text.strip():
                markdown_content += f"{paragraph.text.strip()}\n\n"
        for table in doc.tables:
            markdown_content += FileProcessor.table_to_markdown(table)
    except Exception as e:
        print(f"Error reading Word file: {e}")
    return markdown_content

def handle_image_file(file_path):
    """
    Processes an image file for LLM compatibility.
    Extracts text and metadata, and generates a Markdown report.
    """
    file_path = FileProcessor.clean_path(file_path)
    try:
        with Image.open(file_path) as img:
            metadata = {
                "format": img.format,
                "size": img.size,
                "mode": img.mode
            }
            extracted_text = pytesseract.image_to_string(img)
            markdown_content = f"# Image Metadata\n"
            for key, value in metadata.items():
                markdown_content += f"- **{key.capitalize()}**: {value}\n"
            markdown_content += "\n## Extracted Text\n"
            markdown_content += extracted_text if extracted_text.strip() else "No text found."
            return markdown_content
    except Exception as e:
        print(f"Failed to process image: {e}")
    return ""

def transcribe_audio_file(file_path):
    """
    Transcribes an audio file to text using speech recognition.
    """
    file_path = FileProcessor.clean_path(file_path)
    recognizer = sr.Recognizer()
    _, ext = os.path.splitext(file_path)
    if ext.lower() not in [".wav", ".wave"]:
        print("Only WAV files are supported for now.")
        return None
    try:
        with sr.AudioFile(file_path) as source:
            audio = recognizer.record(source)
            text = recognizer.recognize_google(audio)
            return text
    except sr.UnknownValueError:
        print("Google Speech Recognition could not understand audio.")
    except sr.RequestError as e:
        print(f"Google Speech Recognition error: {e}")
    return None

def table_to_markdown(table):
    """
    Converts a table to Markdown format.
    """
    table_md = ""
    for row in table:
        table_md += "| " + " | ".join(row) + " |\n"
    if table and len(table[0]) > 0:
        table_md += "|---" * len(table[0]) + "|\n"
    return table_md

def transcribe_audio(file_path):
    """
    Transcribes an audio file to text using speech recognition.
    Supports only WAV files for now.
    """
    file_path = FileProcessor.clean_path(file_path)
    recognizer = sr.Recognizer()
    _, ext = os.path.splitext(file_path)

    # Ensure the file is in WAV format
    if ext.lower() not in [".wav", ".wave"]:
        return "Error: Only WAV files are supported for transcription."

    try:
        with sr.AudioFile(file_path) as source:
            audio = recognizer.record(source)
            text = recognizer.recognize_google(audio)
            return text
    except sr.UnknownValueError:
        return "Error: Audio content could not be understood."
    except sr.RequestError as e:
        return f"Error: Speech recognition service request failed: {e}"
    except Exception as e:
        return f"Error: Transcription failed: {e}"

